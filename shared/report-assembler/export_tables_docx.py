#!/usr/bin/env python3
"""
export_tables_docx.py — 三线表 Word 导出器
===========================================
将 Markdown 表格或 CSV 数据输出为符合医学期刊三线表格式的 .docx 文件。

三线表格式规范:
  ┌─────────────────────────────────┐  ← 顶线 (粗 2pt)
  │ 表序 + 表题                      │
  ├─────────────────────────────────┤  ← 表头下线 (中 1pt)
  │ 表头 | 数据                      │
  ├─────────────────────────────────┤
  │ 数据行 (仅大段间隔时加辅助横线)    │
  ├─────────────────────────────────┤  ← 底线 (粗 2pt)
  │ 表注 (如有)                      │
  └─────────────────────────────────┘
  无竖线, 无左右端线

用法:
  python export_tables_docx.py \\
    --input table_data.csv \\
    --output reports/tables/table1.docx \\
    --title "表1 基线特征" \\
    --note "数据以均值±标准差表示"
    
  python export_tables_docx.py \\
    --input-md "|变量|组A|组B|\\n|---|---|---|\\n|Age|...|" \\
    --output reports/tables/table2.docx \\
    --title "表2 Logistic回归结果"

依赖: python-docx
作者: MSRA Team
版本: 0.1.0
"""

import argparse
import csv
import io
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================================
# 三线表格式常量
# ============================================================================

# 边框宽度 (磅)
BORDER_TOP = Pt(2)       # 顶线粗
BORDER_HEADER = Pt(1)    # 表头下线中等
BORDER_BOTTOM = Pt(2)    # 底线粗
BORDER_NONE = Pt(0)      # 无线

# 字体
FONT_NAME = "宋体"
FONT_NAME_HEADER = "黑体" if os.name == "nt" else "Arial"
FONT_SIZE = 10  # 小五号 ≈ 10pt
FONT_SIZE_TITLE = 10  # 五号 ≈ 10.5pt

# 颜色
COLOR_TEXT = RGBColor(0, 0, 0)
COLOR_WHITE = RGBColor(255, 255, 255)


# ============================================================================
# 辅助函数
# ============================================================================

def _set_cell_border(cell, **kwargs):
    """设置单元格的边框。

    kwargs 可以是: top, bottom, left, right, insideH, insideV
    每个值为字典: {'sz': Pt(n), 'color': '000000', 'val': 'single'}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="nil" w:sz="0" w:space="0" w:color="auto"/>'
                          f'<w:left w:val="nil" w:sz="0" w:space="0" w:color="auto"/>'
                          f'<w:bottom w:val="nil" w:sz="0" w:space="0" w:color="auto"/>'
                          f'<w:right w:val="nil" w:sz="0" w:space="0" w:color="auto"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)

    for edge, props in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" '
                f'w:sz="{int(props.get("sz", 0) / 635 * 10000) if isinstance(props.get("sz"), Emu) else int(props.get("sz", 0) * 8)}" '
                f'w:space="0" w:color="{props.get("color", "000000")}"/>'
            )
            tcBorders.append(element)
        else:
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), str(int(props.get("sz", 0) * 8)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), props.get('color', '000000'))

    return cell


def _remove_all_cell_borders(cell):
    """移除单元格的所有边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_elem in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(border_elem)


def _set_run_font(run, name=FONT_NAME, size=FONT_SIZE, bold=False):
    """设置 run 的字体属性。"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = COLOR_TEXT
    # West/East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), name)


def _set_cell_alignment(cell, horizontal=WD_ALIGN_PARAGRAPH.CENTER, vertical='center'):
    """设置单元格对齐。"""
    # 段落对齐
    for paragraph in cell.paragraphs:
        paragraph.alignment = horizontal
        # 垂直居中
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = tcPr.find(qn('w:vAlign'))
        if vAlign is None:
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            tcPr.append(vAlign)
        else:
            vAlign.set(qn('w:val'), 'center')

    # 段落间距
    pf = cell.paragraphs[0].paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)


# ============================================================================
# 三线表核心生成
# ============================================================================

def _apply_three_line_borders(table, n_rows, n_cols, header_rows=1):
    """给表格应用三线表边框。

    三线表只有三条横线:
      - 顶线: 粗 (所有列宽)
      - 表头下线: 中 (所有列宽)
      - 底线: 粗 (所有列宽)
    """
    _THICK = 2   # 2 pt
    _MID = 1     # 1 pt
    _NONE = 0

    # 遍历所有非空行
    for row_idx in range(n_rows):
        for col_idx in range(n_cols):
            cell = table.cell(row_idx, col_idx)
            top = _THICK if row_idx == 0 else _NONE
            bottom = (_MID if row_idx == header_rows - 1
                       else _THICK if row_idx == n_rows - 1
                       else _NONE)

            # 设置边框
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # 移除旧边框
            for old_border in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old_border)

            borders_xml = '<w:tcBorders {}>'.format(nsdecls('w'))
            for edge in ['top', 'left', 'bottom', 'right']:
                if edge == 'top':
                    sz = top
                elif edge == 'bottom':
                    sz = bottom
                else:
                    sz = 0  # 左/右无边框

                if sz > 0:
                    borders_xml += (
                        f'<w:{edge} w:val="single" w:sz="{sz * 8}" '
                        f'w:space="0" w:color="000000"/>'
                    )
                else:
                    borders_xml += (
                        f'<w:{edge} w:val="nil" w:sz="0" '
                        f'w:space="0" w:color="auto"/>'
                    )
            borders_xml += '</w:tcBorders>'

            tcPr.append(parse_xml(borders_xml))


def _set_table_width(table, doc):
    """设置表格宽度为页面正文宽度。"""
    section = doc.sections[0]
    page_width = section.page_width
    left_margin = section.left_margin
    right_margin = section.right_margin
    body_width = page_width - left_margin - right_margin

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')


def md_table_to_docx(doc: Document, md_table: str, title: str = "",
                     note: str = "", header_rows: int = 1) -> Document:
    """将 Markdown 表格写入 docx，应用三线表格式。

    Parameters
    ----------
    doc : Document
        已打开的 docx 文档对象
    md_table : str
        Markdown 表格文本
    title : str
        表序 + 表题（如 "表1 基线特征"）
    note : str
        表注（如 "数据以均值±标准差表示"）
    header_rows : int
        表头行数 (默认 1)

    Returns
    -------
    Document
        docx 文档对象
    """
    # --- 解析 Markdown 表格 ---
    lines = [l.strip() for l in md_table.strip().split('\n') if l.strip()]
    table_lines = [l for l in lines if l.startswith('|')]

    # 过滤分隔行
    def is_sep(line):
        s = line.replace('|', '').replace('-', '').replace(':', '').strip()
        return len(s) == 0

    data_lines = []
    for line in table_lines:
        if is_sep(line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        data_lines.append(cells)

    if not data_lines:
        return doc

    n_cols = max(len(row) for row in data_lines)
    n_rows = len(data_lines)

    # --- 添加表题 ---
    if title:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(title)
        _set_run_font(run, name=FONT_NAME_HEADER, size=FONT_SIZE_TITLE, bold=False)
        p_title.paragraph_format.space_after = Pt(4)
        p_title.paragraph_format.space_before = Pt(6)

    # --- 创建 Word 表格 ---
    table = doc.add_table(rows=n_rows, cols=n_cols)

    # 设置表格内边距
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    # 禁用自动调整
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    # --- 填充数据 ---
    for row_idx, cells in enumerate(data_lines):
        for col_idx, cell_text in enumerate(cells):
            if col_idx >= n_cols:
                break
            cell = table.cell(row_idx, col_idx)

            # 清空默认段落
            cell.text = ""

            # 添加文本
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)

            # 表头加粗
            is_header = row_idx < header_rows
            _set_run_font(run, name=FONT_NAME, size=FONT_SIZE, bold=is_header)

            # 对齐: 第一列左对齐，其他居中
            align = (WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0
                     else WD_ALIGN_PARAGRAPH.CENTER)
            p.alignment = align

            # 垂直居中
            _set_cell_alignment(cell, horizontal=align)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

    # --- 应用三线表边框 ---
    _apply_three_line_borders(table, n_rows, n_cols, header_rows)

    # --- 自动列宽 ---
    auto_fit = tblPr.find(qn('w:tblW'))
    if auto_fit is not None:
        tblPr.remove(auto_fit)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    # --- 表注 ---
    if note:
        p_note = doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_note.add_run(note)
        _set_run_font(run, name=FONT_NAME, size=9, bold=False)
        p_note.paragraph_format.space_before = Pt(4)


def csv_to_docx(doc: Document, csv_path: str, title: str = "",
                note: str = "", header_rows: int = 1) -> Document:
    """将 CSV 文件写入 docx，应用三线表格式。"""
    with open(csv_path, encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return doc

    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    # 表题
    if title:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(title)
        _set_run_font(run, name=FONT_NAME_HEADER, size=FONT_SIZE_TITLE, bold=False)
        p_title.paragraph_format.space_after = Pt(4)
        p_title.paragraph_format.space_before = Pt(6)

    # 表格
    table = doc.add_table(rows=n_rows, cols=n_cols)

    for row_idx, row_data in enumerate(rows):
        for col_idx in range(min(len(row_data), n_cols)):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(row_data[col_idx])

            is_header = row_idx < header_rows
            _set_run_font(run, name=FONT_NAME, size=FONT_SIZE, bold=is_header)

            align = (WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0
                     else WD_ALIGN_PARAGRAPH.CENTER)
            p.alignment = align
            _set_cell_alignment(cell, horizontal=align)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

    # 三线表边框
    _apply_three_line_borders(table, n_rows, n_cols, header_rows)

    # 表注
    if note:
        p_note = doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_note.add_run(note)
        _set_run_font(run, name=FONT_NAME, size=9, bold=False)
        p_note.paragraph_format.space_before = Pt(4)

    return doc


# ============================================================================
# CLI
# ============================================================================

def parse_args():
    parser = argparse.ArgumentParser(
        description="MSRA 三线表 Word 导出器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  从CSV: export_tables_docx.py --input data.csv --output t.docx --title \"表1 基线\"\n"
            "  内联MD: export_tables_docx.py "
            '--input-md "|A|B|\\n|---|---|\\n|1|2|" --output t.docx'
        ))
    parser.add_argument("--input", help="CSV 文件路径")
    parser.add_argument("--input-md", help="Markdown 表格文本 (内联)")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    parser.add_argument("--title", default="", help="表序 + 表题")
    parser.add_argument("--note", default="", help="表注")
    parser.add_argument("--header-rows", type=int, default=1,
                        help="表头行数 (默认 1)")
    return parser.parse_args()


def main():
    args = parse_args()

    doc = Document()

    # 页面设置 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    if args.input:
        csv_to_docx(doc, args.input, args.title, args.note, args.header_rows)
    elif args.input_md:
        md_table_to_docx(doc, args.input_md, args.title, args.note, args.header_rows)
    else:
        print("错误: 必须提供 --input 或 --input-md", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ 三线表已导出: {output_path.resolve()}")


if __name__ == "__main__":
    main()
